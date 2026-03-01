import unittest
import io
from docx import Document
import sys
import os

# Add parent directory to sys.path so we can import from api
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

# We need to mock some env vars before importing api.index
os.environ['SUPABASE_URL'] = 'http://localhost'
os.environ['SUPABASE_KEY'] = 'dummy'
os.environ['BOT_KEY'] = 'dummy'

from api.index import generate_casting_docx

class TestGenerateCastingDocx(unittest.TestCase):
    def test_generate_casting_docx_empty(self):
        project_name = "Test Project"
        applications = []

        result_io = generate_casting_docx(applications, project_name)

        self.assertIsInstance(result_io, io.BytesIO)

        # Load the generated document
        doc = Document(result_io)

        # Check header
        self.assertEqual(len(doc.paragraphs), 2)
        self.assertIn("КАСТИНГ-ЛИСТ: Test Project", doc.paragraphs[0].text)
        self.assertIn("Сформировано:", doc.paragraphs[1].text)

        # Check table
        self.assertEqual(len(doc.tables), 1)
        table = doc.tables[0]
        # Only header rows might be present depending on implementation, but we didn't add rows for apps
        self.assertEqual(len(table.rows), 2) # Init creates 2 rows (range(2))

    def test_generate_casting_docx_with_data(self):
        project_name = "Another Project"
        applications = [
            {
                "full_name": "Ivan Ivanov",
                "city": "Moscow",
                "age": "25",
                "height_weight": "180/75",
                "sizes": "M",
                "fee_range": "1000",
                "phone": "+79991234567",
                "instagram": "@ivan",
                "experience": "Some exp",
                "video_audition_url": "http://video.com",
                "portfolio_url": "http://portfolio.com",
                "photo_urls": "[\"http://photo1.com\", \"http://photo2.com\"]"
            }
        ]

        result_io = generate_casting_docx(applications, project_name)

        self.assertIsInstance(result_io, io.BytesIO)

        doc = Document(result_io)

        # Check title
        self.assertIn("КАСТИНГ-ЛИСТ: Another Project", doc.paragraphs[0].text)

        # Check table
        table = doc.tables[0]
        # Should have 2 init rows + 1 app row = 3 rows
        self.assertEqual(len(table.rows), 3)

        row = table.rows[2] # the app row

        # Left column
        c_info = row.cells[0]
        text = c_info.text

        self.assertIn("Ivan Ivanov", text)
        self.assertIn("Moscow", text)
        self.assertIn("25", text)
        self.assertIn("180/75", text)
        self.assertIn("M", text)
        self.assertIn("1000", text)
        self.assertIn("+79991234567", text)
        self.assertIn("@ivan", text)
        self.assertIn("Some exp", text)
        self.assertIn("http://video.com", text)
        self.assertIn("http://portfolio.com", text)
        self.assertIn("http://photo1.com", text)
        self.assertIn("http://photo2.com", text)

        # Right column
        c_photo = row.cells[1]
        self.assertIn("[Место для фото]", c_photo.text)

    def test_generate_casting_docx_missing_fields(self):
        project_name = "Missing Fields Project"
        applications = [
            {
                # only name
                "full_name": "Anna Karenina",
            }
        ]

        result_io = generate_casting_docx(applications, project_name)
        self.assertIsInstance(result_io, io.BytesIO)

        doc = Document(result_io)
        row = doc.tables[0].rows[2]
        c_info = row.cells[0]
        text = c_info.text

        self.assertIn("Anna Karenina", text)
        self.assertNotIn("📍", text)
        self.assertNotIn("🎂", text)
        self.assertNotIn("📏", text)
        self.assertNotIn("👟", text)
        self.assertNotIn("💰 Бюджет", text)
        self.assertNotIn("Tel:", text)
        self.assertNotIn("Inst:", text)
        self.assertNotIn("💡 Опыт:", text)
        self.assertNotIn("🎬 Видео-визитка:", text)
        self.assertNotIn("📂 Портфолио:", text)
        self.assertNotIn("📸 Фотографии", text)

    def test_generate_casting_docx_photo_urls_parsing(self):
        project_name = "Photo URLs Project"
        applications = [
            {
                "full_name": "Test Photos",
                "photo_urls": '["https://example.com/1.jpg", "https://example.com/2.png"]'
            },
            {
                "full_name": "Test Single String Photo",
                "photo_urls": 'https://example.com/single.jpg'
            },
            {
                "full_name": "Test Invalid Photo",
                "photo_urls": 'not a url'
            }
        ]

        result_io = generate_casting_docx(applications, project_name)
        doc = Document(result_io)

        # Row 2 is the first app
        text1 = doc.tables[0].rows[2].cells[0].text
        self.assertIn("https://example.com/1.jpg", text1)
        self.assertIn("https://example.com/2.png", text1)

        # Row 3 is the second app
        text2 = doc.tables[0].rows[3].cells[0].text
        self.assertIn("https://example.com/single.jpg", text2)

        # Row 4 is the third app
        text3 = doc.tables[0].rows[4].cells[0].text
        self.assertNotIn("not a url", text3)

if __name__ == '__main__':
    unittest.main()
